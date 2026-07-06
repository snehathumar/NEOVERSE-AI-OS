import io
import json
from abc import ABC, abstractmethod
from backend.reporting.models import ExecutiveReport

class BaseExportGenerator(ABC):
    @abstractmethod
    def generate(self, report: ExecutiveReport) -> bytes:
        pass

class MarkdownGenerator(BaseExportGenerator):
    def generate(self, report: ExecutiveReport) -> bytes:
        md = f"# {report.executive_summary}\n\n"
        md += f"## Business Context\n{report.business_context}\n\n"
        md += f"## Verified Evidence\n{report.verified_evidence_summary}\n\n"
        md += f"## Key Assumptions\n{report.key_assumptions}\n\n"
        md += f"## Decision Analysis\n{report.decision_analysis}\n\n"
        md += f"## Expert Debate Summary\n{report.expert_debate_summary}\n\n"
        md += f"## Simulation Summary\n{report.simulation_summary}\n\n"
        md += f"## Recommended Action Plan\n{report.recommended_action_plan}\n\n"
        return md.encode('utf-8')

class JSONGenerator(BaseExportGenerator):
    def generate(self, report: ExecutiveReport) -> bytes:
        return report.model_dump_json(indent=2).encode('utf-8')

class PDFGenerator(BaseExportGenerator):
    def generate(self, report: ExecutiveReport) -> bytes:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.drawString(100, 750, "NEOVERSE AI OS - Executive Report")
        c.drawString(100, 730, f"Session ID: {report.metadata.session_id}")
        
        # Super simple layout for demonstration - would use Platypus for real reports
        textobject = c.beginText(100, 700)
        textobject.textLine("Executive Summary:")
        # Simple text wrapping simulation
        lines = report.executive_summary.split('\n')
        for line in lines[:10]: # truncate for basic layout
            textobject.textLine(line[:80])
            
        c.drawText(textobject)
        c.showPage()
        c.save()
        return buffer.getvalue()

class DOCXGenerator(BaseExportGenerator):
    def generate(self, report: ExecutiveReport) -> bytes:
        from docx import Document
        
        doc = Document()
        doc.add_heading('NEOVERSE AI OS - Executive Report', 0)
        
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(report.executive_summary)
        
        doc.add_heading('Business Context', level=1)
        doc.add_paragraph(report.business_context)
        
        doc.add_heading('Decision Analysis', level=1)
        doc.add_paragraph(report.decision_analysis)
        
        doc.add_heading('Explainability Audit', level=1)
        doc.add_paragraph(report.explainability.why_this_recommendation)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
